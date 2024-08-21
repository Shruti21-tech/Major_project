import streamlit as st
from textblob import TextBlob
import language_tool_python
from io import StringIO
from io import BytesIO
from docx import Document

class SpellCheckerModule:
    def __init__(self):
        self.tool = language_tool_python.LanguageTool('en-US')

    def correct_spell(self, text):
        words = text.split()
        corrected_words = []
        for word in words:
            corrected_word = str(TextBlob(word).correct())
            corrected_words.append(corrected_word)
        return " ".join(corrected_words)

    def correct_grammar(self, text):
        matches = self.tool.check(text)
        corrections = []
        for match in matches:
            corrections.append({
                'incorrect': match.context,
                'suggestions': match.replacements
            })
        corrected_text = language_tool_python.utils.correct(text, matches)
        return corrected_text, corrections, len(matches)

# Streamlit App
st.title("Spell and Grammar Checker")

spell_checker_module = SpellCheckerModule()

# Spell Checker and Grammar Correction Section
st.header("Text Input Checker")
input_text = st.text_area("Enter text for spell and grammar checking:")
if st.button("Check Spelling and Grammar"):
    corrected_text = spell_checker_module.correct_spell(input_text)
    corrected_text, grammar_corrections, grammar_mistakes = spell_checker_module.correct_grammar(corrected_text)
    
    st.write("Corrected Text:")
    st.write(corrected_text)
    
    st.write(f"Grammar Mistakes Found: {grammar_mistakes}")
    for correction in grammar_corrections:
        st.write(f"Incorrect: {correction['incorrect']}")
        st.write(f"Suggestions: {', '.join(correction['suggestions'])}")

# File Upload Section for DOCX Files
st.header("File Upload Checker")
uploaded_file = st.file_uploader("Upload a Word file (.docx) for spell and grammar checking:", type=["docx"])
if uploaded_file is not None:
    # Read the DOCX file
    doc = Document(uploaded_file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    file_content = '\n'.join(full_text)
    
    # Correct spelling and grammar in the file
    corrected_file_text = spell_checker_module.correct_spell(file_content)
    corrected_file_text, corrected_file_grammar, grammar_mistakes = spell_checker_module.correct_grammar(corrected_file_text)
    
    st.write("Corrected File Text:")
    st.write(corrected_file_text)

    st.write(f"Grammar Mistakes Found: {grammar_mistakes}")
    for correction in corrected_file_grammar:
        st.write(f"Incorrect: {correction['incorrect']}")
        st.write(f"Suggestions: {', '.join(correction['suggestions'])}")

    # Optionally, save the corrected text back to a DOCX file
    output_doc = Document()
    output_doc.add_paragraph(corrected_file_text)
    output_io = StringIO()
    output_io = BytesIO()
    output_io.seek(0)
    output_doc.save(output_io)
    
    st.download_button(
        label="Download Corrected File",
        data=output_io.getvalue(),
        file_name="corrected_file.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )